"""Generate BRD as a .docx file for the Hosted Checkout Monitoring System."""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), "Hosted_Checkout_Monitoring_BRD.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(table):
    """Add thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "CCCCCC")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row_obj = table.rows[ri + 1]
        bg = "EAF0FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row_obj.cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    _set_cell_borders(table)
    doc.add_paragraph()  # spacing after table


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.size = Pt(18)


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.size = Pt(14)


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.size = Pt(11)


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = "Normal"
    for run in p.runs:
        run.font.size = Pt(10)


def add_callout(doc: Document, title: str, body: str, color_hex: str = "FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, color_hex)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    for run in p2.runs:
        run.font.size = Pt(9)
    _set_cell_borders(table)
    doc.add_paragraph()


def add_scenario_block(doc: Document, sid: str, name: str, dimension: str, description: str,
                        thresholds: list, high_risk: list, note: str):
    p = doc.add_paragraph()
    run_id = p.add_run(f"[{sid}]  ")
    run_id.bold = True
    run_id.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run_id.font.size = Pt(11)
    run_name = p.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(11)
    run_dim = p.add_run(f"  ({dimension})")
    run_dim.font.size = Pt(9)
    run_dim.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_body(doc, description)

    add_table(doc, ["Parameter", "Default Value"], thresholds, col_widths=[3.2, 2.0])

    if high_risk:
        p2 = doc.add_paragraph()
        r = p2.add_run("High-Risk Sub-Classification Thresholds")
        r.italic = True
        r.font.size = Pt(9)
        add_table(doc, ["Parameter", "Default / Note"], high_risk, col_widths=[3.2, 2.8])

    if note:
        np = doc.add_paragraph()
        nr = np.add_run(f"Note: {note}")
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()


# ── main ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hosted Checkout Monitoring System")
    r2.font.size = Pt(15)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Document Version", "1.0"),
        ("Status", "Draft"),
        ("Prepared", datetime.date.today().strftime("%B %Y")),
        ("Classification", "Internal / Restricted"),
        ("System", "Hosted Checkout Card Cash-In AML Monitoring"),
    ]
    for i, (k, v) in enumerate(fields):
        meta.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        meta.rows[i].cells[1].paragraphs[0].add_run(v)
        for cell in meta.rows[i].cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    _set_cell_borders(meta)
    doc.add_page_break()

    # ── 1. Purpose & Scope ────────────────────────────────────────────────
    add_h1(doc, "1. Purpose & Scope")
    add_body(doc, (
        "This document defines the functional and data requirements for the Hosted Checkout Monitoring "
        "System — a web-based AML surveillance platform that ingests card cash-in transaction data, applies "
        "configurable rule-based detection scenarios, and provides a structured case-management workflow for "
        "compliance investigators and supervisors."
    ))
    add_callout(doc,
        "Key Requirement — Automated Daily Pipeline",
        "Transaction data must be ingested automatically every day by pulling the previous day's file from "
        "the Switch Portal. No manual upload step is required for routine daily operations. Manual upload "
        "capability is retained as a fallback for remediation and historical data loads.",
        "D6E4F0"
    )
    add_body(doc, (
        "The system covers card-not-present (hosted checkout) transactions only. Each card is identified by "
        "the combination of BIN (bank identification number) and the last four digits of the card number, "
        "linked to a mobile wallet identified by its MSISDN. Cash-in events are the sole transaction "
        "direction in scope."
    ))

    # ── 2. Stakeholders & User Roles ──────────────────────────────────────
    add_h1(doc, "2. Stakeholders & User Roles")
    add_table(doc,
        ["Role", "Access Level", "Capabilities"],
        [
            ["admin", "User management only",
             "Create, activate, deactivate, and reset passwords for users. No access to detections or transactions."],
            ["supervisor", "Full operational access",
             "Upload/import batches, run all scenarios, view and manage all detections, force-set any status, export, configure thresholds, manage investigator policy, view transactions explorer."],
            ["investigator", "Triage & notes",
             "View detections, advance status per the investigator policy, add notes. Cannot import, export, run scenarios, configure thresholds, or edit/delete others' notes."],
        ],
        col_widths=[1.2, 1.8, 3.8]
    )
    add_body(doc, (
        "Investigator status transitions are governed by a configurable policy set by the supervisor. "
        "Supervisors may force-set any status regardless of the normal workflow graph."
    ))

    # ── 3. Automated Daily Data Pipeline ─────────────────────────────────
    add_h1(doc, "3. Automated Daily Data Ingestion Pipeline")
    add_body(doc, (
        "The pipeline is fully automated. It runs on a daily schedule, extracts the previous day's "
        "transaction data from the Switch Portal, validates and deduplicates it, stores it in the "
        "application database, and immediately runs all active monitoring scenarios to produce "
        "detections — with no manual intervention required under normal operations."
    ))

    pipeline_steps = [
        ("Step 1 — Scheduled Extract", "Automated scheduler (daily)",
         "A scheduled job runs daily at a configured time and connects to the Switch Portal API (or SFTP "
         "endpoint) to download the previous day's hosted-checkout transaction file in Excel (.xlsx) format."),
        ("Step 2 — File Validation & Parse", "System",
         "The file is validated for required columns. Header normalisation tolerates minor casing/spacing "
         "differences. Malformed files are rejected with a descriptive error logged."),
        ("Step 3 — Deduplication", "System",
         "Rows are deduplicated by UniqueId. Intra-file duplicates keep the first occurrence. Rows with "
         "UniqueIds already present in the database are skipped silently, enabling safe re-delivery."),
        ("Step 4 — Import Batch Creation", "System",
         "All new rows are written to the transaction_rows table under a new import_batch record. The batch "
         "status transitions from 'uploaded' to 'ready'. Skipped-row counts are logged for auditability."),
        ("Step 5 — Scenario Run (All 6)", "System (auto-triggered)",
         "Immediately after a successful import, all enabled scenarios (D1, D2, D3, W1, W2, W3) are run "
         "automatically. Detections are persisted with status = 'new'."),
        ("Step 6 — Wallet Enrichment", "System",
         "If a MariaDB enrichment connection is configured, wallet profiles (holder name, city) are fetched "
         "and merged into detection metrics. Enrichment failures are logged and retryable from the UI."),
        ("Step 7 — Notification & Audit", "System",
         "A pipeline completion event is logged. Any error triggers an alert to configured recipients and "
         "marks the batch as 'failed' for manual retry."),
    ]
    for title, owner, desc in pipeline_steps:
        add_table(doc,
            ["Stage", "Owner"],
            [[title, owner]],
            col_widths=[4.0, 2.8]
        )
        add_body(doc, desc)

    add_h3(doc, "3.1 Switch Portal Integration")
    add_table(doc,
        ["Attribute", "Requirement"],
        [
            ["Source system", "Switch Portal (payment-switch transaction reporting)"],
            ["Trigger", "Scheduled daily job (configurable time, default: 01:00 local)"],
            ["Protocol", "REST API or SFTP — to be confirmed with Switch Portal team"],
            ["File format", "Excel (.xlsx), same column schema as manual import"],
            ["Data window", "Previous calendar day (T-1) by default; configurable look-back for recovery"],
            ["Credentials", "Stored in environment / secrets manager; not in codebase"],
            ["Retry policy", "3 automatic retries with exponential back-off on transient errors"],
            ["Failure handling", "Mark batch as 'failed'; alert compliance operations; allow manual retry from UI"],
            ["Max file size", "Configurable via MAX_UPLOAD_BYTES environment variable"],
        ],
        col_widths=[2.2, 4.6]
    )

    add_h3(doc, "3.2 Required Input Columns")
    add_table(doc,
        ["Column Name", "Mandatory", "Description"],
        [
            ["RequestTimestamp", "Required", "Transaction date/time; used for day and rolling-window grouping"],
            ["Mobile", "Required", "WalletId — the mobile-wallet MSISDN"],
            ["Bin", "Required", "First 6–8 digits of the card number (BIN)"],
            ["AccountNumberLast4", "Required", "Last 4 digits of the card number"],
            ["Credit", "Required", "Transaction amount (cash-in value)"],
            ["ReasonCode", "Required", "0 = approved; non-zero = rejected/failed"],
            ["TransactionId", "Required", "Source system transaction identifier"],
            ["UniqueId", "Required (web)", "Stable row-level deduplication key for web imports"],
            ["AccountHolder", "Optional", "Card holder name; enriches detection metrics"],
            ["OPP_card.issuer.bank", "Optional", "Card issuer bank; used for bank-filter and UniqueBanks metric"],
        ],
        col_widths=[2.0, 1.2, 3.6]
    )

    add_h3(doc, "3.3 Deduplication Rules")
    add_table(doc,
        ["Rule", "Behaviour"],
        [
            ["Intra-file duplicate UniqueId", "Keep first occurrence, skip subsequent rows; count logged"],
            ["UniqueId already in database", "Skip silently; count logged in batch record"],
            ["All rows duplicate / no new rows", "Batch marked failed with message 'No new rows to import'"],
            ["Re-delivery of same file", "Safe — no double-counting; idempotent by design"],
        ],
        col_widths=[3.0, 3.8]
    )

    # ── 4. Monitoring Scenarios ───────────────────────────────────────────
    add_h1(doc, "4. Monitoring Scenarios")
    add_body(doc, (
        "Six scenarios are defined across two cadences. Daily scenarios operate on calendar-day aggregates. "
        "Weekly scenarios use a rolling 7-day sliding window anchored to each individual transaction "
        "timestamp. All thresholds are configurable in the UI without a code deploy. Each scenario can "
        "be independently enabled or disabled."
    ))

    add_h2(doc, "4.1 Daily Scenarios")

    add_scenario_block(doc, "D1", "Many Cards → One Wallet", "Wallet / Day",
        "A single mobile wallet receives approved cash-in loads from many distinct physical cards on the "
        "same calendar day, suggesting a structured funnel where multiple card-holders push value to a "
        "single beneficiary wallet.",
        [
            ["Min amount per transaction", "50,000"],
            ["Min total daily amount", "500,000"],
            ["Min transaction count (approved)", "3"],
            ["Min unique cards", "3"],
        ],
        [
            ["High-risk min total amount", "configurable (default: 0 = off)"],
            ["High-risk min expenditure %", "configurable (default: 0 = off)"],
        ],
        "Approved transactions only count toward TxnCount and TotalAmount. Failed/non-approved attempts "
        "are tracked separately as NotApprovedCount."
    )

    add_scenario_block(doc, "D2", "One Card → Multiple Wallets", "Card / Day",
        "A single physical card (identified by BIN + last-4) is used to load multiple distinct mobile "
        "wallets on the same calendar day, suggesting one card is being used to distribute value across "
        "many accounts.",
        [
            ["Min amount per transaction", "50,000"],
            ["Min total daily amount", "500,000"],
            ["Min unique wallets", "3"],
        ],
        [
            ["High-risk min total amount", "configurable (default: 0 = off)"],
            ["High-risk min wallet expenditure %", "configurable (default: 0 = off)"],
            ["High-risk min wallets %", "configurable (default: 0 = off)"],
        ],
        "CardId = BIN concatenated with AccountNumberLast4."
    )

    add_scenario_block(doc, "D3", "Multiple Failed Transactions", "Wallet / Day",
        "A mobile wallet accumulates many rejected/failed hosted-checkout attempts on the same calendar "
        "day (ReasonCode ≠ 0), indicating probing behavior or card-testing activity.",
        [["Min rejected transaction count", "5"]],
        [],
        "Detection covers all rejected rows regardless of per-transaction amount."
    )

    add_h2(doc, "4.2 Weekly Scenarios (Rolling 7-Day Window)")
    add_body(doc, (
        "Each transaction acts as a potential window end-point. The window covers [T-6 days, T] inclusive. "
        "Approved amounts and card/wallet counts are maintained in a sliding-window deque for performance. "
        "Results are deduplicated to one detection per entity per window end-date."
    ))

    add_scenario_block(doc, "W1", "Many Cards → One Wallet (Rolling 7 Days)", "Wallet / 7-day window",
        "Rolling 7-day sliding window per wallet. Triggers when approved cash-ins from many distinct cards "
        "exceed both the transaction count and total amount thresholds within any 7-day period, catching "
        "patterns that spread across calendar weeks.",
        [
            ["Min transaction count (approved)", "10"],
            ["Min unique cards", "3"],
            ["Min total amount (approved)", "500,000"],
        ],
        [],
        "Window end = date of each individual transaction. Deduplication keeps the first hit per wallet × end-date."
    )

    add_scenario_block(doc, "W2", "One Card → Multiple Wallets (Rolling 7 Days)", "Card / 7-day window",
        "Rolling 7-day sliding window per card. Triggers when one physical card loads many distinct wallets "
        "and the approved total amount crosses the threshold within any 7-day period.",
        [
            ["Min unique wallets", "5"],
            ["Min transaction count (approved)", "1"],
            ["Min total amount (approved)", "500,000"],
        ],
        [],
        "Window mechanics identical to W1 but pivoted on CardId instead of WalletId."
    )

    add_scenario_block(doc, "W3", "Multiple Failed Transactions (Rolling 7 Days)", "Wallet / 7-day window",
        "Rolling 7-day sliding window per wallet tracking rejected attempts only. Triggers when cumulative "
        "rejected count crosses the threshold within any 7-day period.",
        [["Min rejected count", "10"]],
        [],
        "No approved/amount filter — any rejected row (ReasonCode ≠ 0) counts."
    )

    add_h3(doc, "4.3 Bank-Specific Monitoring Filter")
    add_body(doc, (
        "Each scenario independently supports an optional issuer-bank substring filter. When set, the "
        "scenario only considers transactions where the card issuer bank name contains the configured "
        "substring (case-insensitive). Multiple bank substrings can be pipe-separated (e.g. 'BankA|BankB'). "
        "When blank, no bank filter is applied and all issuers are included."
    ))

    add_h3(doc, "4.4 High-Risk Classification (D1 & D2)")
    add_table(doc,
        ["Scenario", "Parameter", "Meaning"],
        [
            ["D1", "d1_risk_min_total_amount", "Min total daily load amount for wallet to be flagged 'high risk'"],
            ["D1", "d1_risk_min_expenditure_pct", "Min observed debit % of total wallet expenditure to flag 'high risk'"],
            ["D2", "d2_risk_min_total_amount", "Min total daily load amount for card to be flagged 'high risk'"],
            ["D2", "d2_risk_min_wallet_expenditure_pct", "Min max single-wallet expenditure % to flag 'high risk'"],
            ["D2", "d2_risk_min_wallets_pct", "Min % of matched wallets relative to total wallets to flag 'high risk'"],
        ],
        col_widths=[0.8, 2.8, 3.2]
    )
    add_body(doc, (
        "High-risk sub-fields default to 0 (disabled). Setting any value above 0 enables high-risk "
        "sub-classification, surfaced as a 'Risk: high' label on qualifying detections and filterable "
        "from the detections list."
    ))

    # ── 5. Detection Metrics ──────────────────────────────────────────────
    add_h1(doc, "5. Detection Metrics Captured")
    add_body(doc, "Each detection stores a JSON metrics snapshot at creation time. The following fields are captured per scenario where applicable:")
    add_table(doc,
        ["Metric Field", "Label", "Scenarios"],
        [
            ["WalletId", "MSISDN (wallet)", "D1, D3, W1, W3"],
            ["CardId", "Card ID (BIN + last-4)", "D2, W2"],
            ["TxnDate", "Detection date", "D1, D2, D3"],
            ["TxnWeek", "Rolling window end date", "W1, W2, W3"],
            ["TxnCount", "Approved transaction count", "D1, D2, W1, W2"],
            ["UniqueCards", "Distinct cards involved", "D1, W1"],
            ["UniqueWallets", "Distinct wallets involved", "D2, W2"],
            ["TotalAmount", "Sum of approved cash-in amounts", "D1, D2, W1, W2"],
            ["AvgAmount", "Average transaction amount", "All"],
            ["MinAmount", "Minimum transaction amount", "All"],
            ["MaxAmount", "Maximum transaction amount", "All"],
            ["RejectedCount", "Count of rejected transactions", "D3, W3"],
            ["NotApprovedCount", "Count of non-approved attempts", "All"],
            ["CardHolderNamesPipe", "Pipe-delimited card holder names", "All"],
            ["UniqueBanks", "Distinct issuer banks", "All (if bank column present)"],
            ["TopCardId", "Card with highest total amount", "D1, W1"],
            ["TopCardTotalAmount", "Top card's total load amount", "D1, W1"],
            ["TopCardHolderName", "Top card holder name(s)", "D1, W1"],
            ["WalletHolderName / WalletHolderFullName", "Wallet holder name (enriched)", "D1, D3, W1, W3"],
            ["WalletCityName", "Wallet residency city (enriched)", "D1, D3, W1, W3"],
            ["WalletHolderNamesPipe", "Pipe-delimited wallet holder names", "D2, W2"],
            ["Risk", "High / Low risk classification", "D1, D2"],
            ["RiskObservedExpenditurePct", "Observed debit percentage", "D1"],
            ["RiskObservedWalletsPct", "Observed wallets percentage", "D2"],
        ],
        col_widths=[2.6, 2.4, 1.8]
    )

    # ── 6. Detection Workflow ─────────────────────────────────────────────
    add_h1(doc, "6. Detection Case Management Workflow")
    add_body(doc, (
        "Detections move through a defined status lifecycle. Supervisors may force any transition; "
        "investigators are restricted to transitions permitted by the configurable investigator policy. "
        "All status changes are logged to a persistent status_history table with actor name and timestamp."
    ))

    add_h3(doc, "6.1 Status Definitions")
    add_table(doc,
        ["Status", "Key", "Description"],
        [
            ["Test", "test", "Created by a supervisor test run; not a live alert."],
            ["New", "new", "Live detection awaiting first triage."],
            ["False Positive (Initial)", "false_positive_initial", "Investigator's initial assessment: low risk."],
            ["Suspicious (Initial)", "suspicious_initial", "Investigator's initial assessment: warrants deeper review."],
            ["Pending Evidence", "pending_evidence", "Waiting for additional documentation or data."],
            ["False Positive (Final)", "false_positive_final", "Closed — supervisor confirms no risk."],
            ["Suspicious (Final)", "suspicious_final", "Closed — confirmed suspicious activity."],
            ["Wallet Lock", "wallet_lock", "Terminal — wallet has been locked."],
            ["Wallet CI", "wallet_ci", "Terminal — case escalated to Criminal Investigations."],
        ],
        col_widths=[1.8, 2.2, 2.8]
    )

    add_h3(doc, "6.2 Allowed Status Transitions")
    add_table(doc,
        ["From Status", "Allowed Next Statuses"],
        [
            ["test", "new"],
            ["new", "false_positive_initial, suspicious_initial, pending_evidence"],
            ["false_positive_initial", "false_positive_final, suspicious_final, pending_evidence, wallet_lock, wallet_ci"],
            ["suspicious_initial", "false_positive_final, suspicious_final, pending_evidence, wallet_lock, wallet_ci"],
            ["false_positive_final", "wallet_lock, wallet_ci, pending_evidence"],
            ["suspicious_final", "wallet_lock, wallet_ci, pending_evidence"],
            ["pending_evidence", "new, false_positive_initial, suspicious_initial"],
            ["wallet_lock", "(terminal — no further transitions)"],
            ["wallet_ci", "(terminal — no further transitions)"],
        ],
        col_widths=[2.2, 4.6]
    )
    add_body(doc, (
        "Supervisors bypass the transition graph and may force-set any status. Investigator transitions "
        "are further constrained by the per-deployment investigator policy (allowed_map)."
    ))

    add_h3(doc, "6.3 Notes System")
    add_table(doc,
        ["Capability", "Supervisor", "Investigator"],
        [
            ["Add note to detection", "Yes", "Yes"],
            ["Edit own note", "Yes", "No"],
            ["Delete note", "Yes", "No"],
            ["View all notes", "Yes", "Yes"],
        ],
        col_widths=[3.2, 1.4, 1.4]
    )

    add_h3(doc, "6.4 Prior Detection Cross-Reference")
    add_body(doc, (
        "The detection detail page automatically queries for prior detections linked to the same wallet "
        "(MSISDN) tokens, surfacing historical activity to assist investigators in identifying repeat "
        "patterns across multiple import batches."
    ))

    # ── 7. Transactions Explorer ──────────────────────────────────────────
    add_h1(doc, "7. Transactions Explorer")
    add_body(doc, (
        "Supervisors can search across all imported transaction batches. Results are paginated "
        "and ordered by timestamp descending."
    ))
    add_table(doc,
        ["Filter", "Type", "Notes"],
        [
            ["Batch ID", "Integer", "Scope results to a single import batch; omit for all batches"],
            ["MSISDN (wallet)", "Substring", "Case-insensitive partial match on WalletId"],
            ["Card ID", "Substring", "Case-insensitive partial match on CardId (BIN + last-4)"],
            ["Account Holder", "Substring", "Case-insensitive partial match on AccountHolder name"],
            ["Issuer Bank", "Substring", "Case-insensitive partial match on OPP_card.issuer.bank"],
            ["Amount Min / Max", "Numeric", "Inclusive range filter on transaction amount"],
            ["Date From / To", "Date + Time", "Filter by RequestTimestamp / TxnTimestamp"],
            ["Approved", "Boolean", "'true' or 'false'"],
            ["Per page", "Integer", "20 / 50 / 100 / 200"],
        ],
        col_widths=[1.8, 1.2, 3.8]
    )

    # ── 8. Scenario Configuration ─────────────────────────────────────────
    add_h1(doc, "8. Scenario Configuration & Thresholds")
    add_body(doc, (
        "A singleton scenario_config database row holds all threshold values and enable/disable flags. "
        "The Scenarios Manager page provides a per-scenario UI. All changes take effect immediately on "
        "the next scenario run — no restart required."
    ))

    add_h3(doc, "8.1 Complete Threshold Reference")
    add_table(doc,
        ["DB Column", "Scenario(s)", "Default", "Description"],
        [
            ["d_amount_min", "D1", "50,000", "Min per-transaction amount (D1: wallet/day)"],
            ["d_total_amount_min", "D1", "500,000", "Min total amount per group per day (D1: wallet/day)"],
            ["d1_min_txn", "D1", "3", "Min approved transaction count (wallet/day)"],
            ["d1_min_unique_cards", "D1", "3", "Min distinct cards (wallet/day)"],
            ["d1_risk_min_total_amount", "D1", "0 (off)", "High-risk total amount threshold"],
            ["d1_risk_min_expenditure_pct", "D1", "0 (off)", "High-risk expenditure % threshold"],
            ["d2_amount_min", "D2", "50,000", "Min per-transaction amount (D2: card/day)"],
            ["d2_total_amount_min", "D2", "500,000", "Min total amount per group per day (D2: card/day)"],
            ["d2_min_wallets", "D2", "3", "Min distinct wallets (card/day)"],
            ["d2_risk_min_total_amount", "D2", "0 (off)", "High-risk total amount threshold"],
            ["d2_risk_min_wallet_expenditure_pct", "D2", "0 (off)", "High-risk per-wallet expenditure % threshold"],
            ["d2_risk_min_wallets_pct", "D2", "0 (off)", "High-risk matched-wallets % threshold"],
            ["d3_min_rejected", "D3", "5", "Min rejected attempts (wallet/day)"],
            ["w1_min_txn", "W1", "10", "Min approved txn count in rolling 7-day window"],
            ["w1_min_unique_cards", "W1", "3", "Min distinct cards in rolling 7-day window"],
            ["w1_min_total_amount", "W1", "500,000", "Min approved total amount in rolling 7-day window"],
            ["w2_min_wallets", "W2", "5", "Min distinct wallets in rolling 7-day window"],
            ["w2_min_txn", "W2", "1", "Min approved txn count in rolling 7-day window"],
            ["w2_min_total_amount", "W2", "500,000", "Min approved total amount in rolling 7-day window"],
            ["w3_min_rejected", "W3", "10", "Min rejected count in rolling 7-day window"],
            ["monitored_banks (per scenario)", "All", "null", "Optional issuer bank substring filter (pipe-separated)"],
            ["scenario_enabled (per scenario)", "All", "true", "Enable / disable each scenario independently"],
        ],
        col_widths=[2.2, 1.0, 1.0, 2.6]
    )

    add_h3(doc, "8.2 Scenario Test-Run Capability")
    add_body(doc, (
        "Before applying new thresholds to production, a supervisor may run a single scenario in 'test' "
        "mode against any ready import batch. Test detections are created with status = 'test' and can "
        "be reviewed and bulk-deleted without affecting live detections."
    ))

    # ── 9. Wallet Enrichment ──────────────────────────────────────────────
    add_h1(doc, "9. Wallet Enrichment (External Lookup)")
    add_body(doc, (
        "When a MariaDB/MySQL enrichment database is configured, the system fetches wallet holder "
        "profiles (full name and city) and merges them into detection metrics at scenario-run time."
    ))
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Enrichment source", "MariaDB / MySQL — actors_clean1_clone table (msisdn, Fullname, city)"],
            ["Configuration", "MINITRANS_HOST / PORT / USER / PASSWORD / DATABASE environment variables"],
            ["City code mapping", "Optional GOV_MAPPING_PATH Excel file: city code → display name"],
            ["Failure behaviour", "Scenarios still run; enrichment fields left empty; logged for retry"],
            ["Manual retry", "Supervisor can trigger enrichment retry from the Scenarios Manager page"],
            ["Metrics enriched", "WalletHolderName, WalletHolderFullName, WalletCityName, WalletHolderNamesPipe, WalletCityNamesPipe"],
        ],
        col_widths=[2.0, 4.8]
    )

    # ── 10. Reporting & Export ────────────────────────────────────────────
    add_h1(doc, "10. Reporting & Export")
    add_table(doc,
        ["Output", "Format", "Who", "Notes"],
        [
            ["Filtered detections export", "Excel (.xlsx)", "Supervisor", "Respects all active filters (status, scenario, batch, date range, MSISDN, risk)"],
            ["Desktop scenario outputs", "Excel (.xlsx per scenario)", "Desktop tool", "One file per scenario: Scenario_D1_daily.xlsx … Scenario_W3_weekly.xlsx"],
            ["Import batch history", "In-app table", "Supervisor", "Filename, row count, status, creation date"],
            ["Detection detail", "HTML", "Supervisor / Investigator", "Metrics snapshot, linked transactions, notes, status history"],
            ["Status history audit trail", "In-app + DB", "Supervisor", "Full log of transitions with actor name and timestamp"],
        ],
        col_widths=[1.8, 1.4, 1.4, 2.2]
    )

    # ── 11. Non-Functional Requirements ───────────────────────────────────
    add_h1(doc, "11. Non-Functional Requirements")
    add_table(doc,
        ["Category", "Requirement"],
        [
            ["Deployment", "Podman Compose (RHEL 9) — PostgreSQL + FastAPI web service; single-command startup"],
            ["Authentication", "Session-cookie (signed); configurable timeout; HTTPS required in production"],
            ["Password security", "Bcrypt-hashed passwords; password change enforceable on first login"],
            ["Database migrations", "Alembic; auto-applied at container startup (alembic upgrade head)"],
            ["File upload limit", "Configurable via MAX_UPLOAD_BYTES environment variable"],
            ["Audit trail", "All detection status changes logged with actor name and timestamp; immutable"],
            ["Health endpoint", "GET /health returns {ok: true, database: true}; returns 503 if DB unreachable"],
            ["Scheduler reliability", "Daily pipeline job must be idempotent; safe to re-run for same date without duplication"],
            ["Error alerting", "Pipeline failures must alert compliance operations within 15 minutes"],
            ["Observability", "Structured logs for import events, scenario runs, enrichment, and pipeline activity"],
        ],
        col_widths=[2.0, 4.8]
    )

    # ── 12. Assumptions & Dependencies ────────────────────────────────────
    add_h1(doc, "12. Assumptions & Dependencies")
    add_table(doc,
        ["Item", "Detail"],
        [
            ["Switch Portal API / SFTP", "Switch Portal must expose a daily export endpoint with the required column schema. Credentials and endpoint URL are provided by the Switch Portal team."],
            ["Time zone", "Transaction timestamps are stored and compared in the application's configured time zone. All daily aggregations use calendar days in that zone."],
            ["Card identification", "Cards are uniquely identified by BIN + AccountNumberLast4. This key may match multiple physical cards sharing the same combination — a known limitation of last-4 identification."],
            ["Wallet = MSISDN", "WalletId is the mobile subscriber number (MSISDN). One wallet per MSISDN is assumed."],
            ["Enrichment DB optional", "The MariaDB enrichment connection is optional; all scenarios run without it, with enrichment fields blank."],
            ["Single deployment", "This BRD covers a single-tenant deployment. Multi-tenancy is out of scope."],
        ],
        col_widths=[2.0, 4.8]
    )

    # ── 13. Glossary ──────────────────────────────────────────────────────
    add_h1(doc, "13. Glossary")
    add_table(doc,
        ["Term", "Definition"],
        [
            ["BIN", "Bank Identification Number — first 6–8 digits of a payment card, identifying the issuing bank and card scheme."],
            ["CardId", "System composite key: BIN concatenated with AccountNumberLast4 (e.g. '4111111234')."],
            ["Cash-in / Load", "A credit event where value is added to a mobile wallet via a card payment."],
            ["Hosted Checkout", "Card-not-present online payment flow where the cardholder enters card details in a hosted payment page."],
            ["MSISDN", "Mobile Station International Subscriber Directory Number — the mobile phone number used as WalletId."],
            ["ReasonCode", "Response code from the payment switch; 0 = approved, non-zero = rejected/failed."],
            ["Rolling window", "A sliding 7-calendar-day lookback period anchored to each individual transaction timestamp."],
            ["Scenario", "A rule-based detection algorithm that examines grouped transaction data and produces detections when thresholds are exceeded."],
            ["Switch Portal", "The payment-switch reporting portal that provides raw hosted-checkout transaction data for daily ingestion."],
            ["UniqueId", "A stable per-row identifier in the source file used for deduplication across imports."],
            ["WalletId", "The MSISDN identifying the mobile wallet that received the cash-in."],
        ],
        col_widths=[1.8, 5.0]
    )

    # Footer paragraph
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Hosted Checkout Monitoring System — BRD v1.0  |  Confidential — Internal Use Only")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
